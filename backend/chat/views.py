# chat/views.py

from rest_framework import viewsets, status
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.utils import timezone
from datetime import timedelta
from django.http import JsonResponse, FileResponse
from django.views.decorators.csrf import ensure_csrf_cookie
from docx import Document
from docx.shared import Pt
import io, json, logging, re
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from django.shortcuts import get_object_or_404
from .models import AiPlan
from openai import OpenAI
from django.conf import settings

from .models import ChatSession, ChatMessage, Fon, AiPlan, UserFeedback, AgentLog
from .serializers import (
    ChatSessionSerializer,
    ChatMessageSerializer,
    ChatSessionListSerializer
)
from agent_ai.agent_ai_service import AgentAIService

logger = logging.getLogger(__name__)
agent_service = AgentAIService()


@ensure_csrf_cookie
def get_csrf_token(request):
    return JsonResponse({"success": True})


class ChatSessionViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]
    serializer_class = ChatSessionSerializer

    def get_queryset(self):
        return ChatSession.objects.filter(user=self.request.user)

    def get_serializer_class(self):
        if self.action == 'list':
            return ChatSessionListSerializer
        return ChatSessionSerializer

    def perform_create(self, serializer):
        ai_model = self.request.data.get('ai_model', 'openai')
        if ai_model not in ['openai', 'lstm']:
            ai_model = 'openai'
            logger.warning("Geçersiz ai_model, 'openai' kullanıldı")
        serializer.save(user=self.request.user, ai_model=ai_model)

    def update(self, request, *args, **kwargs):
        instance = self.get_object()
        ai_model = request.data.get('ai_model')
        if ai_model in ['openai', 'lstm'] and ai_model != instance.ai_model:
            instance.ai_model = ai_model
            instance.save(update_fields=['ai_model', 'updated_at'])
            logger.info(f"Model güncellendi: {instance.id} -> {ai_model}")
        return super().update(request, *args, **kwargs)

    @action(detail=True, methods=["get"], url_path='messages')
    def messages(self, request, pk=None):
        session = self.get_object()
        qs = ChatMessage.objects.filter(session=session).order_by('timestamp')
        serializer = ChatMessageSerializer(qs, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=["post"])
    def send_message(self, request, pk=None):
        session = self.get_object()
        full_prompt = request.data.get("message", "").strip()
        raw_input = request.data.get("raw_input")

        if not full_prompt:
            return Response({"error": "message boş olamaz"}, status=status.HTTP_400_BAD_REQUEST)

        content_to_store = raw_input if raw_input is not None else full_prompt
        user_msg = ChatMessage.objects.create(session=session, is_user=True, content=content_to_store)

        ai_resp = _get_ai_response(full_prompt, session.ai_model)
        print(f"AI yanıtı türü: {type(ai_resp)}")
        if isinstance(ai_resp, dict):
            print(f"plan_id: {ai_resp.get('plan_id')}")
    
        ai_content = ai_resp["content"] if isinstance(ai_resp, dict) else ai_resp
        ai_msg = ChatMessage.objects.create(session=session, is_user=False, content=ai_content)

    # Plan ilişkisini kur
        if isinstance(ai_resp, dict) and ai_resp.get("plan_id") is not None:
            plan_id = ai_resp.get("plan_id")
            print(f"Plan ID: {plan_id} ile ilişkilendiriliyor")
        
            try:
                plan = AiPlan.objects.get(id=plan_id)
                session.plan = plan  # ForeignKey ilişkisi
                session.save(update_fields=["plan", "updated_at"])
                print(f"Plan {plan.id} session {session.id} ile ilişkilendirildi")
            except AiPlan.DoesNotExist:
                print(f"Plan bulunamadı: {plan_id}")

        data = {
            "ai_response": ai_content,
            "user_message_id": user_msg.id,
            "ai_message_id": ai_msg.id,
        }
        if isinstance(ai_resp, dict):
            data.update(meta=ai_resp.get("meta"), plan_id=ai_resp.get("plan_id"))

        return Response(data)

    @action(detail=True, methods=['get'], url_path='generate_pdf')
    def generate_pdf(self, request, pk=None):
        session = self.get_object()
        plan = getattr(session, 'plan', None)
        if not plan:
            return Response({'error': 'Plan yok.'}, status=404)

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=72)
        styles = getSampleStyleSheet()
        elems = [
            Paragraph(f"TÜBİTAK {plan.fon.kod} – {plan.fon.tur} Planı", styles['Heading1']),
            Spacer(1, 0.2 * inch)
        ]
        for p in plan.plan_metni.split('\n\n'):
            elems += [Paragraph(p, styles['Normal']), Spacer(1, 0.1 * inch)]
        doc.build(elems)
        buf.seek(0)
        return FileResponse(buf, as_attachment=True, filename=f"plan_{plan.id}.pdf")

    
    



def _get_ai_response(message: str, model_type: str):
    try:
        m = model_type.lower().strip()
        
        # LSTM model için
        if m == "lstm":
            resp, meta = agent_service.generate_response(message, None, None)
            content = resp if isinstance(resp, str) else resp.get("content", "")
            content += "\n\n[TÜBİTAK LSTM AI tarafından oluşturuldu]"
            
            # Plan içeriği kontrolü
            plan_id = None
            if "Ay 1:" in content:
                print("Plan içeriği tespit edildi!")
                # Fon kodunu mesajdan çıkarmaya çalış
                import re
                fon_code = None
                fon_match = re.search(r'(2209-[AB]|2247-C|2205|TÜBİTAK\s+\w+)', message, re.IGNORECASE)
                if fon_match:
                    fon_code = fon_match.group(1).strip()
                
                try:
                    # Fon kaydını bul veya varsayılan kullan
                    fon = None
                    if fon_code:
                        try:
                            fon = Fon.objects.get(kod__icontains=fon_code.split()[0])
                        except Fon.DoesNotExist:
                            pass
                    
                    if not fon:
                        fon = Fon.objects.first()
                        
                    if not fon:
                        # Hiç fon yoksa oluştur
                        fon = Fon.objects.create(
                            kod="TÜBİTAK",
                            tur="Genel",
                            ay_suresi=12
                        )
                    
                    # Ay sayısını belirle
                    ay_suresi = 12  # Varsayılan
                    try:
                        # Metinden en yüksek ay numarasını bul
                        ay_pattern = r'Ay\s+(\d+):'
                        ay_matches = re.findall(ay_pattern, content)
                        if ay_matches:
                            ay_numbers = [int(m) for m in ay_matches]
                            ay_suresi = max(ay_numbers)
                    except:
                        pass
                    
                    # Plan oluştur
                    plan = AiPlan.objects.create(
                        fon=fon,
                        ay_suresi=ay_suresi,
                        plan_metni=content
                    )
                    plan_id = plan.id
                    print(f"Yeni plan oluşturuldu: {plan_id}")
                except Exception as e:
                    print(f"Plan oluşturma hatası: {e}")
            
            return {"content": content, "meta": meta, "plan_id": plan_id}
        
        # OpenAI için
        api_key = settings.OPENAI_API_KEY
        if not api_key:
            return {"content": "OpenAI API anahtarı eksik.", "meta": None, "plan_id": None}

        client = OpenAI(api_key=api_key)
        out = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Sen bir TÜBİTAK proje asistanısın."},
                {"role": "user", "content": message}
            ],
            temperature=0.7,
            max_tokens=1500
        )
        
        content = out.choices[0].message.content
        
        # Plan içeriği kontrolü (OpenAI için de aynı)
        plan_id = None
        if "Ay 1:" in content:
            print("Plan içeriği tespit edildi!")
            import re
            fon_code = None
            fon_match = re.search(r'(2209-[AB]|2247-C|2205|TÜBİTAK\s+\w+)', message, re.IGNORECASE)
            if fon_match:
                fon_code = fon_match.group(1).strip()
            
            try:
                # Fon kaydını bul veya varsayılan kullan
                fon = None
                if fon_code:
                    try:
                        fon = Fon.objects.get(kod__icontains=fon_code.split()[0])
                    except Fon.DoesNotExist:
                        pass
                
                if not fon:
                    fon = Fon.objects.first()
                    
                if not fon:
                    # Hiç fon yoksa oluştur
                    fon = Fon.objects.create(
                        kod="TÜBİTAK",
                        tur="Genel",
                        ay_suresi=12
                    )
                
                # Ay sayısını belirle
                ay_suresi = 12  # Varsayılan
                try:
                    # Metinden en yüksek ay numarasını bul
                    ay_pattern = r'Ay\s+(\d+):'
                    ay_matches = re.findall(ay_pattern, content)
                    if ay_matches:
                        ay_numbers = [int(m) for m in ay_matches]
                        ay_suresi = max(ay_numbers)
                except:
                    pass
                
                # Plan oluştur
                plan = AiPlan.objects.create(
                    fon=fon,
                    ay_suresi=ay_suresi,
                    plan_metni=content
                )
                plan_id = plan.id
                print(f"Yeni plan oluşturuldu: {plan_id}")
            except Exception as e:
                print(f"Plan oluşturma hatası: {e}")
        
        return {"content": content, "meta": None, "plan_id": plan_id}

    except Exception as e:
        logger.error("AI yanıt hatası", exc_info=True)
        return {"content": f"Hata: {str(e)[:200]}", "meta": None, "plan_id": None}


def _render_plan_pdf_for_session(session: ChatSession):
    plan = getattr(session, 'plan', None)
    if not plan:
        return Response({'error': 'Plan yok'}, status=status.HTTP_404_NOT_FOUND)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    elems = [
        Paragraph(f"TÜBİTAK {plan.fon.kod} – {plan.fon.tur} Planı", styles['Heading1']),
        Spacer(1, 0.2 * inch),
    ]
    for p in plan.plan_metni.split('\n\n'):
        elems.append(Paragraph(p, styles['Normal']))
        elems.append(Spacer(1, 0.1 * inch))

    table_data = [["Ay", "Aktivite"]]
    acts = re.findall(r'Ay\s*(\d+):\s*(.+)', plan.plan_metni)
    if acts:
        for ay, txt in acts:
            table_data.append([ay, txt.strip()])
        tbl = Table(table_data, colWidths=[0.7*inch, 4*inch])
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elems.append(Spacer(1, 0.2 * inch))
        elems.append(tbl)

    doc.build(elems)
    buf.seek(0)
    filename = f"Tubitak_{plan.fon.kod}_Plan_{plan.id}.pdf"
    return FileResponse(buf, as_attachment=True, filename=filename)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def cleanup_chat_sessions(request):
    cutoff = timezone.now() - timedelta(hours=8)
    ChatSession.objects.filter(user=request.user, updated_at__lt=cutoff, is_active=True).update(is_active=False)
    return Response(status=status.HTTP_204_NO_CONTENT)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def fon_listesi(request):
    data = [
        {'id': f.id, 'kod': f.kod, 'tur': f.tur, 'ay_suresi': f.ay_suresi, 'aciklama': f.aciklama}
        for f in Fon.objects.all().order_by('kod')
    ]
    return Response(data)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def lstm_predict(request):
    try:
        fon = Fon.objects.get(id=request.data.get('fon_id'))
        query = request.data.get('query') or f"{fon.kod} fonu için {fon.ay_suresi} aylık plan"
        resp, meta = agent_service.generate_response(query, fon.id, {
            'kod': fon.kod, 'tur': fon.tur, 'ay_suresi': fon.ay_suresi
        })
        plan = AiPlan.objects.create(fon=fon, ay_suresi=fon.ay_suresi,
                                     plan_metni=resp, meta_data=json.dumps(meta, default=str))
        AgentLog.objects.create(islem_tipi='tahmin',
                                detay=f"{request.user.username}", plan=plan)
        return Response({'success': True, 'plan_id': plan.id, 'plan': resp, 'meta': meta})
    except Fon.DoesNotExist:
        return Response({'success': False, 'error': 'Fon bulunamadı'}, status=status.HTTP_404_NOT_FOUND)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def plan_feedback(request):
    try:
        plan = AiPlan.objects.get(id=request.data.get('plan_id'))
        score = int(request.data.get('score'))
        plan.feedback_score = score
        plan.save()
        UserFeedback.objects.create(plan=plan, user=request.user,
                                    puan=score, yorum=request.data.get('comment', ''))
        AgentLog.objects.create(islem_tipi='geri_bildirim',
                                detay=f"{request.user.username}", plan=plan)
        return Response({'success': True})
    except AiPlan.DoesNotExist:
        return Response({'success': False, 'error': 'Plan bulunamadı'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        logger.error("plan_feedback hata", exc_info=True)
        return Response({'success': False, 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def feedback(request):
    try:
        data = request.data
        for f in ('query', 'response', 'score'):
            if f not in data:
                return Response({'error': f"{f} gerekli"}, status=status.HTTP_400_BAD_REQUEST)
        ok = agent_service.process_feedback(
            query=data['query'],
            response=data['response'],
            score=data['score'],
            comment=data.get('comment', ''),
            features=data.get('features', {})
        )
        return Response({'message': ok and 'Kaydedildi' or 'Hata'},
                        status= status.HTTP_201_CREATED if ok else status.HTTP_500_INTERNAL_SERVER_ERROR)
    except Exception as e:
        logger.error("feedback hata", exc_info=True)
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def test_model(request):
    """
    Anonim da erişebilir. ?model_type=openai|lstm&message=... 
    """
    model_type = request.query_params.get('model_type', 'openai')
    msg = request.query_params.get('message', 'Merhaba, test.')
    resp = _get_ai_response(msg, model_type)
    if isinstance(resp, dict):
        return Response({
            'model_type': model_type,
            'response': resp.get('content'),
            'is_lstm': model_type == 'lstm',
            'meta': resp.get('meta')
        })
    else:
        return Response({
            'model_type': model_type,
            'response': resp,
            'is_lstm': False
        })


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def generate_plan_doc(request, pk):
    """
    Seans içindeki planı .docx olarak indirir.
    """
    try:
        # Oturumu kullanıcıya göre al
        session = get_object_or_404(ChatSession, pk=pk, user=request.user)
        print(f"Session bulundu: {session.id}")
        
        # Plan kontrolü
        if session.plan is None:
            print(f"Plan ilişkisi yok, plan mesajı aranıyor...")
            
            # Plan içeren bir mesaj var mı kontrol et
            plan_messages = session.messages.filter(is_user=False, content__contains="Ay 1:")
            if plan_messages.exists():
                print(f"Plan içeren mesaj bulundu")
                # Fonu bul veya oluştur
                fon = Fon.objects.first()
                if not fon:
                    fon = Fon.objects.create(
                        kod="TÜBİTAK",
                        tur="Genel",
                        ay_suresi=12
                    )
                
                # Plan oluştur
                plan_content = plan_messages.first().content
                plan = AiPlan.objects.create(
                    fon=fon,
                    ay_suresi=12,
                    plan_metni=plan_content
                )
                
                # Session ile ilişkilendir
                session.plan = plan
                session.save(update_fields=["plan"])
                print(f"Acil durum planı oluşturuldu: {plan.id}")
            else:
                print(f"Bu oturumda plan bulunamadı: {session.id}")
                return Response({'error': 'Bu seansta plan bulunamadı'}, status=404)
        
        # Plan nesnesini al
        plan = session.plan
        print(f"Plan: {plan.id}")
        
        # --- DOCX oluştur ---
        doc = Document()
        styles = doc.styles['Normal'].font
        styles.name = 'Calibri'
        styles.size = Pt(11)

        doc.add_heading(f"TÜBİTAK {plan.fon.kod} -- {plan.fon.tur} Planı", level=1)
        doc.add_paragraph(f"Oluşturulma modeli: {session.ai_model.upper()}")
        
        # Türkçe tarih formatıyla planın oluşturulma tarihini ekle
        doc.add_paragraph(f"Tarih: {plan.olusturma_tarihi.strftime('%d.%m.%Y %H:%M')}")

        for parag in plan.plan_metni.split('\n\n'):
            doc.add_paragraph(parag)

        # Geçici dosya
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f"plan_{plan.id}.docx"
        print(f"DOC dosyası oluşturuldu: {filename}")
        return FileResponse(buf, as_attachment=True, filename=filename)
    
    except Exception as e:
        print(f"DOC oluşturma hatası: {e}")
        import traceback
        traceback.print_exc()
        return Response({'error': f'DOC oluşturma hatası: {str(e)}'}, status=500)